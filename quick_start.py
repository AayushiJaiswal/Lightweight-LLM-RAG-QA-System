#!/usr/bin/env python3
"""
Quick Start Runner for ERP RAG System
Complete automation of setup, indexing, and launch
"""

import sys
import subprocess
import time
from pathlib import Path


def check_ollama():
    """Check if Ollama server is running"""
    import requests

    try:
        r = requests.get("http://localhost:11434/api/tags", timeout=2)
        return r.status_code == 200
    except Exception:
        return False


def start_ollama():
    """Start Ollama server if not running"""
    if check_ollama():
        print("✓ Ollama is running")
        return True

    print("Starting Ollama server...")
    try:
        subprocess.Popen(["ollama", "serve"], shell=True)
        time.sleep(5)
        if check_ollama():
            print("✓ Ollama started")
            return True
    except Exception as e:
        print(f"❌ Could not start Ollama: {e}")
    print("Run manually: ollama serve")
    return False


def check_model(model: str = "llama3.2"):
    """Check if Ollama model is installed"""
    try:
        result = subprocess.run(
            ["ollama", "list"],
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True,
            shell=True,
        )
        return model in result.stdout
    except Exception as e:
        print(f"Error checking model: {e}")
        return False


def create_samples():
    """Create sample ERP documents (PDF, DOCX, TXT)"""
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter
    import docx

    Path("data/raw").mkdir(parents=True, exist_ok=True)

    # PDF sample
    pdf_file = Path("data/raw/purchase_order_guide.pdf")
    c = canvas.Canvas(str(pdf_file), pagesize=letter)
    c.setFont("Helvetica-Bold", 16)
    c.drawString(50, 750, "PURCHASE ORDER CREATION GUIDE")
    c.setFont("Helvetica", 12)
    y = 720
    lines = [
        "HOW TO CREATE A PURCHASE ORDER:",
        "1. Navigate to Procurement > Purchase Orders > New PO",
        "2. Select vendor from vendor master list",
        "3. Add line items with material codes and quantities",
        "4. Enter delivery date and shipping address",
        "5. Submit for approval",
    ]
    for line in lines:
        c.drawString(50, y, line)
        y -= 20
    c.save()
    print("✓ PDF sample created")

    # DOCX sample
    docx_file = Path("data/raw/expense_reimbursement_policy.docx")
    doc = docx.Document()
    doc.add_heading("Expense Reimbursement Policy", 0)
    doc.add_paragraph(
        "Submit expense reports within 30 days. Receipts required for expenses over $25."
    )
    doc.save(str(docx_file))
    print("✓ DOCX sample created")

    # TXT sample
    txt_file = Path("data/raw/vendor_management_guide.txt")
    txt_file.write_text(
        "VENDOR MANAGEMENT PROCEDURES\n"
        "1. Add new vendor\n"
        "2. Assign category\n"
        "3. Upload documents\n"
    )
    print("✓ TXT sample created")


def run_indexing():
    """Index documents"""
    from config import Config
    from ingestion import DocumentIngester, DocumentChunker
    from embeddings_store import RAGRetriever

    Config.create_directories()

    print("Step 1: Ingesting documents")
    ingester = DocumentIngester(str(Config.RAW_DATA_DIR))
    docs = ingester.ingest_all()
    if not docs:
        print("❌ No documents found. Add documents to data/raw and try again.")
        return False

    print("Step 2: Chunking documents")
    chunker = DocumentChunker(Config.CHUNK_SIZE, Config.CHUNK_OVERLAP)
    chunks = chunker.chunk_documents(docs)

    print("Step 3: Building FAISS index")
    retriever = RAGRetriever(Config.EMBEDDING_MODEL)
    retriever.index_documents(chunks)
    retriever.save(str(Config.EMBEDDINGS_DIR))

    print(f"✓ Indexing complete ({len(docs)} docs, {len(chunks)} chunks)")
    return True


def run_demo():
    """Launch UI"""
    from config import Config
    from embeddings_store import RAGRetriever
    from llm_generation import RAGPipeline
    from ui import RAGInterface

    if not start_ollama():
        return False

    if not check_model(Config.OLLAMA_MODEL):
        print(
            f"❌ Model {Config.OLLAMA_MODEL} not found. Pull with: "
            f"ollama pull {Config.OLLAMA_MODEL}"
        )
        return False

    print("Loading vector store...")
    retriever = RAGRetriever(Config.EMBEDDING_MODEL)
    try:
        retriever.load(str(Config.EMBEDDINGS_DIR))
    except Exception:
        print("❌ Index not found. Run index first.")
        return False

    pipeline = RAGPipeline(retriever, Config.OLLAMA_BASE_URL, Config.OLLAMA_MODEL)
    interface = RAGInterface(pipeline)
    interface.launch(Config.UI_PORT)
    return True


def main():
    if len(sys.argv) < 2:
        print("Usage: python quick_start.py [setup|index|demo|all]")
        return

    cmd = sys.argv[1].lower()

    try:
        if cmd == "setup":
            create_samples()
        elif cmd == "index":
            run_indexing()
        elif cmd == "demo":
            run_demo()
        elif cmd == "all":
            # Agar tumhe sample docs nahi chahiye to create_samples() ko comment kar sakti ho
            create_samples()
            if run_indexing():
                run_demo()
        else:
            print(f"Unknown command: {cmd}")
    except Exception as e:
        print(f"❌ Error: {e}")
        import traceback

        traceback.print_exc()


if __name__ == "__main__":
    main()
